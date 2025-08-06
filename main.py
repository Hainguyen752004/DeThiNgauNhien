import docx
import random
import re
import copy

def read_questions_by_part(path):
    doc = docx.Document(path)
    paragraphs = doc.paragraphs

    part1, part2 = [], []
    current_q = []

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.lower().startswith("câu ") and current_q:
            if is_multiple_choice(current_q):
                part1.append(copy.deepcopy(current_q))
            elif is_true_false(current_q):
                part2.append(copy.deepcopy(current_q))
            current_q = [para]
        else:
            current_q.append(para)

    if current_q:
        if is_multiple_choice(current_q):
            part1.append(copy.deepcopy(current_q))
        elif is_true_false(current_q):
            part2.append(copy.deepcopy(current_q))

    return part1, part2

def is_multiple_choice(q):
    return any(p.text.strip().startswith(("A.", "B.", "C.", "D.")) for p in q[1:])

def is_true_false(q):
    return any(p.text.strip().startswith(("a)", "b)", "c)", "d)")) for p in q[1:])

def clean_choice_text(choice):
    return choice.strip()

def get_label(index, kind="ABC"):
    labels = "ABCDEFGHIJKLMNOPQRSTUVWXYZ" if kind == "ABC" else "abcdefghijklmnopqrstuvwxyz"
    return labels[index] if index < len(labels) else f"({index})"

def get_choices_with_correct_marks(paragraphs, kind="ABC"):
    choices = []
    correct_choices = []

    for p in paragraphs:
        raw = p.text.strip()
        if not raw:
            continue
        choice_text = clean_choice_text(raw)
        is_correct = any(run.bold for run in p.runs)
        choices.append(choice_text)
        if is_correct:
            correct_choices.append(choice_text)
    return choices, correct_choices

def shuffle_choices_and_get_answers(paragraphs, kind="ABC"):
    choices, correct_texts = get_choices_with_correct_marks(paragraphs, kind)
    if not correct_texts:
        return [p.text.strip() for p in paragraphs], []

    shuffled = copy.deepcopy(choices)
    random.shuffle(shuffled)

    correct_labels = []
    for i, choice in enumerate(shuffled):
        if choice in correct_texts:
            correct_labels.append(get_label(i, kind))

    return shuffled, correct_labels

def shuffle_question_block_keep_header(block, kind="ABC"):
    header = block[0].text.strip()
    body = block[1:]
    shuffled_choices, answers = shuffle_choices_and_get_answers(body, kind)
    new_block = [header] + shuffled_choices
    return new_block, answers

def shuffle_and_format_with_renumbering(part, kind="ABC"):
    shuffled = random.sample(part, len(part))
    formatted = []
    answers = []

    for new_index, q in enumerate(shuffled, 1):
        q_text, answer_labels = shuffle_question_block_keep_header(q, kind)

        old_header = q_text[0]
        new_header = f"Câu {new_index}."
        rest_of_text = re.sub(r"^Câu\s*\d+\.*\s*", "", old_header)
        q_text[0] = f"{new_header} {rest_of_text.strip()}"

        formatted.append(q_text)

        if answer_labels:
            answers.append(f"Câu {new_index}: {', '.join(answer_labels)}")
        else:
            answers.append(f"Câu {new_index}: [Không xác định]")

    return formatted, answers

def write_to_docx_with_answers(part1, part2, answers1, answers2, filename):
    doc = docx.Document()
    doc.add_heading("ĐỀ THI TRẮC NGHIỆM", level=0)

    doc.add_heading("PHẦN 1: Trắc nghiệm nhiều lựa chọn", level=1)
    for q in part1:
        for line in q:
            doc.add_paragraph(line)
        doc.add_paragraph("")

    doc.add_heading("PHẦN 2: Trắc nghiệm đúng/sai", level=1)
    for q in part2:
        for line in q:
            doc.add_paragraph(line)
        doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading("ĐÁP ÁN", level=0)

    doc.add_heading("PHẦN 1:", level=1)
    for a in answers1:
        doc.add_paragraph(a)

    doc.add_heading("PHẦN 2:", level=1)
    for a in answers2:
        doc.add_paragraph(a)

    doc.save(filename)

def generate_versions(path, num_versions=4):
    part1, part2 = read_questions_by_part(path)

    for i in range(1, num_versions + 1):
        p1, ans1 = shuffle_and_format_with_renumbering(part1, kind="ABC")
        p2, ans2 = shuffle_and_format_with_renumbering(part2, kind="abc")

        write_to_docx_with_answers(p1, p2, ans1, ans2, f"de_va_dapan_{i}.docx")

# === GỌI CHẠY ===
generate_versions("đê chưa trộn.docx", num_versions=4)
